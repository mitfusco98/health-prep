"""
Utility functions for generating patient prep sheet documents
"""

import os
import io
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from models import (
    Patient,
    Condition,
    Screening,
    Visit,
    Vital,
    LabResult,
    ImagingStudy,
    ConsultReport,
    HospitalSummary,
    Immunization,
)


def generate_prep_sheet_doc(
    patient,
    conditions=None,
    screenings=None,
    vitals=None,
    labs=None,
    imaging=None,
    consults=None,
    hospital=None,
    immunizations=None,
    past_appointments=None,
    last_visit_date=None,
):
    """
    Generate a Word document prep sheet with the patient's information

    Args:
        patient: The patient object
        conditions: List of condition objects
        screenings: List of screening objects
        vitals: List of vital sign objects
        labs: List of lab result objects
        imaging: List of imaging study objects
        consults: List of consult report objects
        hospital: List of hospital summary objects
        immunizations: List of immunization objects
        last_visit_date: Date of the last visit

    Returns:
        bytes: The generated Word document as bytes
    """
    # Create a new Word document
    doc = Document()

    # Title
    doc.add_heading("Patient Preparation Sheet", level=1)

    # Current date section
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y')}")
    # Past Appointments
    if past_appointments:
        doc.add_paragraph(
            f"Last visit: {past_appointments[0].appointment_date.strftime('%m/%d/%Y')}"
        )
        doc.add_heading("Recent Appointments", level=3)
        for appointment in past_appointments:
            status_str = f" ({appointment.status})" if appointment.status else ""
            time_str = (
                appointment.appointment_time.strftime("%I:%M %p")
                if appointment.appointment_time
                else ""
            )
            notes_str = f" - {appointment.notes}" if appointment.notes else ""
            doc.add_paragraph(
                f"• {appointment.appointment_date.strftime('%m/%d/%Y')} {time_str}{status_str}{notes_str}"
            )
    elif last_visit_date:
        doc.add_paragraph(f"Last visit: {last_visit_date.strftime('%B %d, %Y')}")

    # Patient Information
    doc.add_heading("Patient Information", level=2)
    doc.add_paragraph(f"Name: {patient.full_name}")
    doc.add_paragraph(
        f"Date of Birth: {patient.date_of_birth.strftime('%m/%d/%Y')} ({patient.age} years)"
    )
    doc.add_paragraph(f"MRN: {patient.mrn}")
    doc.add_paragraph(f"Sex: {patient.sex}")
    doc.add_paragraph(f"Contact: {patient.phone} / {patient.email}")
    doc.add_paragraph(f"Insurance: {patient.insurance}")

    # Active Conditions
    doc.add_heading("Active Medical Conditions", level=2)
    if conditions:
        active_conditions = [c for c in conditions if c.is_active]
        if active_conditions:
            for condition in active_conditions:
                diagnosed_date = (
                    condition.diagnosed_date.strftime("%m/%d/%Y")
                    if condition.diagnosed_date
                    else "Unknown"
                )
                doc.add_paragraph(f"• {condition.name} (Diagnosed: {diagnosed_date})")
        else:
            doc.add_paragraph("No active conditions recorded.")
    else:
        doc.add_paragraph("No conditions recorded.")

    # Most Recent Vitals
    doc.add_heading("Most Recent Vital Signs", level=2)
    if vitals and len(vitals) > 0:
        latest_vital = vitals[0]  # Assuming vitals are sorted with most recent first
        vital_table = doc.add_table(rows=1, cols=2)
        vital_table.style = "Table Grid"

        hdr_cells = vital_table.rows[0].cells
        hdr_cells[0].text = "Vital Sign"
        hdr_cells[1].text = "Value"

        vitals_data = [
            ("Date", latest_vital.date.strftime("%m/%d/%Y")),
            (
                "Height",
                (
                    f"{(latest_vital.height / 2.54):.1f} in"
                    if latest_vital.height
                    else "Not recorded"
                ),
            ),
            (
                "Weight",
                (
                    f"{(latest_vital.weight * 2.20462):.1f} lbs"
                    if latest_vital.weight
                    else "Not recorded"
                ),
            ),
            ("BMI", f"{latest_vital.bmi:.1f}" if latest_vital.bmi else "Not recorded"),
            (
                "Temperature",
                (
                    f"{(latest_vital.temperature * 9/5 + 32):.1f} °F"
                    if latest_vital.temperature
                    else "Not recorded"
                ),
            ),
            (
                "Blood Pressure",
                (
                    f"{latest_vital.blood_pressure_systolic}/{latest_vital.blood_pressure_diastolic} mmHg"
                    if latest_vital.blood_pressure_systolic
                    else "Not recorded"
                ),
            ),
            (
                "Pulse",
                f"{latest_vital.pulse} bpm" if latest_vital.pulse else "Not recorded",
            ),
            (
                "Respiratory Rate",
                (
                    f"{latest_vital.respiratory_rate} breaths/min"
                    if latest_vital.respiratory_rate
                    else "Not recorded"
                ),
            ),
            (
                "Oxygen Saturation",
                (
                    f"{latest_vital.oxygen_saturation}%"
                    if latest_vital.oxygen_saturation
                    else "Not recorded"
                ),
            ),
        ]

        for vital_name, vital_value in vitals_data:
            row_cells = vital_table.add_row().cells
            row_cells[0].text = vital_name
            row_cells[1].text = str(vital_value)
    else:
        doc.add_paragraph("No vital signs recorded.")

    # Screening Checklist
    doc.add_heading("Screening Checklist", level=2)

    # Define standard screenings based on patient demographics
    standard_screenings = [
        "Annual Physical Exam",
        "Vaccination History",
        "Lipid Panel",
        "Glucose/A1c",
    ]

    # Add sex-specific or age-specific screenings
    if patient.sex.lower() == "female":
        standard_screenings.extend(["Pap Smear", "Mammogram", "DEXA Scan"])

    if patient.age >= 45:
        standard_screenings.append("Colonoscopy")

    if conditions:
        # Check for diabetes and add related screenings
        has_diabetes = any(
            "diabetes" in c.name.lower() for c in conditions if c.is_active
        )
        if has_diabetes:
            standard_screenings.extend(
                ["Diabetic Eye Exam", "Microalbumin", "Foot Exam"]
            )

    # Create a table for the screenings
    screening_table = doc.add_table(rows=1, cols=3)
    screening_table.style = "Table Grid"

    hdr_cells = screening_table.rows[0].cells
    hdr_cells[0].text = "Screening"
    hdr_cells[1].text = "Status"
    hdr_cells[2].text = "Last Done"

    # Add screening status to the table
    for screen_name in standard_screenings:
        row_cells = screening_table.add_row().cells
        row_cells[0].text = screen_name

        # Check if we have this screening in the database
        if screenings:
            matching_screenings = [
                s for s in screenings if s.screening_type.lower() in screen_name.lower()
            ]
            if matching_screenings:
                latest_screening = max(matching_screenings, key=lambda s: s.date)
                row_cells[1].text = "Completed" if latest_screening.completed else "Due"
                row_cells[2].text = (
                    latest_screening.date.strftime("%m/%d/%Y")
                    if latest_screening.date
                    else "N/A"
                )
            else:
                row_cells[1].text = "Not documented"
                row_cells[2].text = "N/A"
        else:
            row_cells[1].text = "Not documented"
            row_cells[2].text = "N/A"

    # Recent Labs
    doc.add_heading("Recent Laboratory Results", level=2)
    if labs and len(labs) > 0:
        lab_table = doc.add_table(rows=1, cols=3)
        lab_table.style = "Table Grid"

        hdr_cells = lab_table.rows[0].cells
        hdr_cells[0].text = "Test Name"
        hdr_cells[1].text = "Result"
        hdr_cells[2].text = "Date"

        for lab in labs[:5]:  # Limit to 5 most recent labs
            row_cells = lab_table.add_row().cells
            row_cells[0].text = lab.test_name
            row_cells[1].text = (
                f"{lab.result} {lab.units}" if lab.units else str(lab.result)
            )
            row_cells[2].text = (
                lab.test_date.strftime("%m/%d/%Y") if lab.test_date else "N/A"
            )
    else:
        doc.add_paragraph("No recent laboratory results.")

    # Recent Imaging
    doc.add_heading("Recent Imaging Studies", level=2)
    if imaging and len(imaging) > 0:
        for study in imaging[:3]:  # Limit to 3 most recent studies
            p = doc.add_paragraph(f"• {study.study_type}: {study.description}")
            p.add_run(f" ({study.study_date.strftime('%m/%d/%Y')})")
            if study.findings:
                doc.add_paragraph(
                    f"  Findings: {study.findings}", style="List Bullet 2"
                )
            if study.impression:
                doc.add_paragraph(
                    f"  Impression: {study.impression}", style="List Bullet 2"
                )
    else:
        doc.add_paragraph("No recent imaging studies.")

    # Recent Consults
    doc.add_heading("Recent Consults/Referrals", level=2)
    if consults and len(consults) > 0:
        for consult in consults[:3]:  # Limit to 3 most recent consults
            p = doc.add_paragraph(f"• {consult.provider}: {consult.specialty}")
            p.add_run(f" ({consult.report_date.strftime('%m/%d/%Y')})")
            if consult.reason:
                doc.add_paragraph(f"  Reason: {consult.reason}", style="List Bullet 2")
            if consult.recommendations:
                doc.add_paragraph(
                    f"  Recommendations: {consult.recommendations}",
                    style="List Bullet 2",
                )
    else:
        doc.add_paragraph("No recent consults/referrals.")

    # Recent Hospital Visits
    doc.add_heading("Recent Hospital Visits", level=2)
    if hospital and len(hospital) > 0:
        for visit in hospital[:3]:  # Limit to 3 most recent hospital visits
            p = doc.add_paragraph(f"• {visit.facility}: {visit.visit_type}")
            admission_date = (
                visit.admission_date.strftime("%m/%d/%Y")
                if visit.admission_date
                else "Unknown"
            )
            discharge_date = (
                visit.discharge_date.strftime("%m/%d/%Y")
                if visit.discharge_date
                else "Unknown"
            )
            p.add_run(f" (Admitted: {admission_date}, Discharged: {discharge_date})")
            if visit.diagnosis:
                doc.add_paragraph(
                    f"  Diagnosis: {visit.diagnosis}", style="List Bullet 2"
                )
            if visit.discharge_summary:
                doc.add_paragraph(
                    f"  Summary: {visit.discharge_summary}", style="List Bullet 2"
                )
    else:
        doc.add_paragraph("No recent hospital visits.")

    # Immunizations
    doc.add_heading("Immunization History", level=2)
    if immunizations and len(immunizations) > 0:
        immunization_table = doc.add_table(rows=1, cols=3)
        immunization_table.style = "Table Grid"

        hdr_cells = immunization_table.rows[0].cells
        hdr_cells[0].text = "Vaccine"
        hdr_cells[1].text = "Date"
        hdr_cells[2].text = "Manufacturer"

        for immunization in immunizations:
            row_cells = immunization_table.add_row().cells
            row_cells[0].text = immunization.vaccine_name
            row_cells[1].text = (
                immunization.administration_date.strftime("%m/%d/%Y")
                if immunization.administration_date
                else "N/A"
            )
            row_cells[2].text = immunization.manufacturer or "N/A"
    else:
        doc.add_paragraph("No immunization records.")

    # Notes section
    doc.add_heading("Provider Notes", level=2)
    for _ in range(5):
        doc.add_paragraph("____________________________________________________")

    # Create a binary stream to save the document to
    f = io.BytesIO()
    doc.save(f)
    f.seek(0)

    return f.read()
